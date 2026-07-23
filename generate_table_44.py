from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


# --- helpers for styling (border + shading) ---

def set_cell_shading(cell, fill_hex: str):
    """fill_hex example: 'EFEFEF' """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_borders(cell, color="B7BCC2", size=6):
    """
    size is in eighths of a point in Word XML.
    size=6 => 0.75pt-ish (thin). color is hex without '#'.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right'):
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tcBorders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), str(size))
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)

def set_run_font(run, name="Calibri", size_pt=9, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


# --- data: Table X (Section 4.4) ---
rows = [
    ("response_id", "Stimulus & identifiers",
     "Unique identifier for each prompt–response pair, encoding the base prompt and condition (e.g., B6-POL-UA)."),
    ("base_id", "Stimulus & identifiers",
     "Identifier of the underlying content-stable base prompt (B1–B8)."),
    ("B1–B8", "Stimulus & identifiers",
     "Base prompt set (content held constant across tone × certainty variants)."),
    ("task_type", "Stimulus & identifiers",
     "Task family of the base prompt (e.g., EXPL, CORR; include ADVI/PROC if used)."),
    ("tone", "Prompt conditions",
     "Prompt tone condition: POL (polite/formal), NEU (neutral), CON (confrontational/pressuring)."),
    ("POL / NEU / CON", "Prompt conditions",
     "Tone levels: POL = polite/formal; NEU = neutral; CON = confrontational/pressuring."),
    ("certainty", "Prompt conditions",
     "Certainty-pressure condition: UA (uncertainty allowed) vs HC (high certainty demand)."),
    ("UA / HC", "Prompt conditions",
     "Certainty levels: UA = uncertainty allowed; HC = high certainty demanded."),
    ("EXPL / CORR", "Task types",
     "Task-family labels (add ADVI/PROC if used)."),
    ("run_order", "Randomization",
     "Randomized ordering assigned to stimuli to reduce order effects."),
    ("run_index", "Randomization",
     "Fixed submission position (1–48) used during data collection after sorting by run_order."),
    ("BP (wrapper) / Body", "Segmentation",
     "Outputs are segmented into a boilerplate wrapper and a body region for analysis."),
    ("BP_open_present", "Segmentation",
     "0/1 indicator: standardized opening wrapper present."),
    ("BP_close_present", "Segmentation",
     "0/1 indicator: standardized closing wrapper present."),
    ("BP_count", "Segmentation",
     "Wrapper components present (0–2): BP_open_present + BP_close_present."),
    ("Body_words", "Segmentation",
     "Word count of the body region; denominator for body-based densities."),
    ("response_words (if used)", "Segmentation",
     "Word count of the full response (wrapper + body)."),
    ("ICP / ICP_density_per1000", "Politeness (body)",
     "In-content politeness strategies in the body; density per 1,000 body words."),
    ("EHI / EHI_density_per1000", "Stance (body)",
     "Epistemic hedging (reduced commitment) in the body; density per 1,000 body words."),
    ("COI / COI_density_per1000", "Stance (body)",
     "Conditionality/options structuring in the body; density per 1,000 body words."),
    ("n-grams / collocations", "Phrase analysis (H3)",
     "Phrase-level measures used to detect recurrent routines and co-occurrence patterns."),
    ("concordance / keyness", "Phrase analysis (H3)",
     "Context inspection and distributional salience (e.g., via log-likelihood) where used."),
    ("n", "Table notation",
     "Number of responses contributing to a cell/mean in a table."),
    ("M", "Table notation (added)",
     "Mean (arithmetic average)."),
    ("f_A, f_B", "Table notation (added)",
     "Token/phrase frequency in corpus A / corpus B (as defined in the keyness comparison)."),
    ("Words_A, Words_B", "Table notation (added)",
     "Total word counts (corpus sizes) for corpus A / corpus B used to compute keyness."),
    ("LL (G²)", "Table notation (added)",
     "Log-likelihood keyness statistic (G-squared) reported for phrase/keyness comparisons."),
]

headers = ("Code / Variable", "Category", "Meaning / Definition")


def main(output_path="Table_X_Section_4_4.docx"):
    doc = Document()

    # Optional caption line (you can delete if you don't want it)
    cap = doc.add_paragraph("Table X. Variable Definitions and Coding Legend for Section 4.4 (updated)")
    cap.runs[0].bold = True

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Set column widths (22% / 22% / 56% roughly on A4/standard page)
    # Word isn't perfect with absolute widths, but this gets very close.
    col_widths = (Cm(4.0), Cm(4.0), Cm(10.2))

    # Header row formatting
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size_pt=9, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr_cells[i], "EFEFEF")
        set_cell_borders(hdr_cells[i], color="B7BCC2", size=6)

    # Body rows
    for code, category, meaning in rows:
        row_cells = table.add_row().cells
        values = (code, category, meaning)

        for i, val in enumerate(values):
            row_cells[i].width = col_widths[i]
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            # Make the "added" codes bold to stand out
            bold_code = (i == 0 and code in {"M", "f_A, f_B", "Words_A, Words_B", "LL (G²)"})
            set_run_font(run, size_pt=9, bold=bold_code)
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_borders(row_cells[i], color="B7BCC2", size=6)

        # Reduce spacing inside cells (compact look)
        for cell in row_cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0

    # Note below table (optional)
    note = doc.add_paragraph("Note. This legend consolidates abbreviations and variable names used throughout this section.")
    for r in note.runs:
        set_run_font(r, size_pt=9, bold=False)

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()